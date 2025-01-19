from django.http import HttpResponse
from django.conf import settings
from django.utils import timezone
from django.db.models import Q

from docx.shared import Inches
from docx import Document
from pathlib import Path
import onnxruntime as ort
import numpy as np
import onnx
import string
import cv2 
import logging

from profiles import models as profiles_model
from app import models as app_model

from operate.camera import Camera
from operate.excepts import *
from operate import (
    model_loader as mload,
    image_handler as imhand,
    face_detector as facedet,
	embedding_generator as emb_gen
)
from operate.excepts import *

# Dont really need
from matplotlib import pyplot as plt

DJANGO_SETTINGS 	= settings
OPERATE_SETTINGS 	= app_model.Setting

logger = logging.getLogger(__name__)


def get_full_name(profile:object, m_initial:bool=True) -> str:
	full_name = profile.l_name.title()

	if profile.f_name:
		full_name += f", {profile.f_name.title()}"

	if profile.m_name and m_initial:
		full_name += f" {profile.m_name[0].upper()}."
	elif profile.m_name:
		full_name += f" {profile.m_name.title()}"

	if profile.suffix:
		full_name += f" {profile.suffix}"

	return full_name

def format_text(text:str) -> str:
	text    = text.strip()
	symbols = string.punctuation.replace('_', ' ')
	translation_table = str.maketrans(symbols, '_' * len(symbols))

	text = text.replace(",", "")
	text = text.translate(translation_table)
	return text

def generate_docx(template_path:str, image_path:str, fields:list[str], data:list[str]):
	doc = Document(template_path)

	# inserts the image
	for p in doc.paragraphs:
		if "[[raw_image]]" in p.text:
			p.clear()
			run = p.add_run()
			run.add_picture(
				image_path, 
				width=Inches(2), 
				height=Inches(2)
			)
			p.alignment = 1

	# fills in the table
	for table in doc.tables:
		for row in table.rows:
			cell = row.cells[1]
			for paragraph in cell.paragraphs:
				for field, value in zip(fields, data):
					if field in paragraph.text:
						for run in paragraph.runs:
							if field in run.text:
								run.text = run.text.replace(field, str(value))
								run.font.name = 'Arial'

	file_name	= f"{format_text(data[-1])}.docx"
	if "[[rank]]" in fields:
		file_name	= f"{format_text(data[-5])}_{file_name}"

	save_path	= settings.MEDIA_ROOT.joinpath(file_name)
	doc.save(save_path)

	return [file_name, save_path]

def save_file(file_name:str, save_path:Path):
	with open(save_path, 'rb') as f:
		response = HttpResponse(
			f.read(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = f'attachment; filename="{file_name}"'

	save_path.unlink()

	return response
