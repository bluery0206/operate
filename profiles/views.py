from django.http import HttpResponse
from django.conf import settings
from django.db.models import Q
from django.shortcuts import (
	render, 
	redirect, 
	get_object_or_404
)	

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx import Document

from datetime import datetime
from pathlib import Path
import os

from .models import Personnel, Inmate, Template
from home.utils import save_profile_picture, get_full_name
from .forms import (
	CreatePersonnel, 
	CreateInmate, 
	UpdatePersonnel, 
	UpdateInmate, 
	TemplateUploadForm
)

cwd_path = Path().cwd()
media_path = cwd_path.joinpath("media")
outputs_path = media_path.joinpath("outputs")

if not outputs_path.exists():
	outputs_path.mkdir(exist_ok=False) 

def personnels(request):
	sort_choices = [
		['l_name', "Last Name"],
		['f_name', "First Name"],
		['age', "Age"],
		['date_profiled',"Date Profiled"],
		['date_assigned',"Date Assigned"],
		['date_relieved',"Date Relieved"],
	]

	order_choices = [
		['descending',"Descending"],
		['ascending', "Ascending"],
	]

	context = {
		'personnels'	: Personnel.objects.exclude(archivepersonnel__isnull=False),
		'ranks'			: [rank[0] for rank in Personnel.RANKS],
		'sort_choices'	: sort_choices,
		'order_choices'	: order_choices,
		'filters'		: {}
	}

	if request.method == "GET":
		reset_filter		= request.GET.get("reset_filter", None)
		reset_search		= request.GET.get("reset_search", None)
		search		= request.GET.get("search", "").strip()

		if not reset_search and search:
			context['personnels'] = Personnel.objects.filter(
		        Q(f_name__icontains=search) |
		        Q(l_name__icontains=search) |
		        Q(m_name__icontains=search)
	        )
		else:
			search = ""

		if reset_filter:
			designation	= None
			rank		= None
			sort_by		= None
			sort_order	= None
		else:
			designation	= request.GET.get("designation", "")
			rank		= request.GET.get("rank", "")
			sort_by		= request.GET.get("sort_by", "")
			sort_order	= request.GET.get("sort_order", "descending")

			designation = designation.strip()

			if designation:
				context['personnels'] = context['personnels'].filter(designation__icontains=designation)

			if rank:
				context['personnels'] = context['personnels'].filter(rank=rank)

			if sort_by:
				if sort_order == "descending":
					sort_by = "-" + sort_by
				context['personnels'] = context['personnels'].order_by(sort_by)

			context['filters'] = {
				'designation': designation,
				'rank': rank,
				'sort_by': sort_by,
				'sort_order': sort_order,
				'sort_order': sort_order,
			}

		context['filters'].update({"search": search})

	return render(request, "profiles/personnels.html", context)


def inmates(request):
	sort_choices = [
		['l_name', "Last Name"],
		['f_name', "First Name"],
		['age', "Age"],
		['date_profiled',"Date Profiled"],
		['date_arresetd',"Date Arrested"],
		['date_committed',"Date Committed"],
	]

	order_choices = [
		['descending',"Descending"],
		['ascending', "Ascending"],
	]

	context = {
		'inmates'	: Inmate.objects.exclude(archiveinmate__isnull=False),
		'sort_choices'	: sort_choices,
		'order_choices'	: order_choices,
		'filters'		: {}
	}

	if request.method == "GET":
		reset_filter		= request.GET.get("reset_filter", "")
		reset_search		= request.GET.get("reset_search", "")
		search		= request.GET.get("search", "").strip()

		if not reset_search and search:
			context['inmates'] = Inmate.objects.filter(
		        Q(f_name__icontains=search) |
		        Q(l_name__icontains=search) |
		        Q(m_name__icontains=search)
	        )
		else:
			search = ""

		if reset_filter:
			crime_violated	= None
			sort_by		= None
			sort_order	= None
		else:
			crime_violated	= request.GET.get("crime_violated", "")
			sort_by		= request.GET.get("sort_by", "")
			sort_order	= request.GET.get("sort_order", "descending")

			crime_violated = crime_violated.strip()

			if crime_violated:
				context['inmates'] = context['inmates'].filter(crime_violated__icontains=crime_violated)

			if sort_by:
				if sort_order == "descending":
					sort_by = "-" + sort_by
				context['inmates'] = context['inmates'].order_by(sort_by)

			context['filters'] = {
				'crime_violated': crime_violated,
				'sort_by': sort_by,
				'sort_order': sort_order,
			}

		context['filters'].update({"search": search})
	return render(request, "profiles/inmates.html", context)

 
def profile_personnel(request, pk):
	context = {
		'profile': get_object_or_404(Personnel, pk=pk),
	}
	return render(request, "profiles/profile_personnel.html", context)
 

def profile_inmate(request, pk):
	context = {
		'profile': get_object_or_404(Inmate, pk=pk),
	}
	return render(request, "profiles/profile_inmate.html", context)


def profile_personnel_update(request, pk):
	profile = get_object_or_404(Personnel, pk=pk)

	if request.method == "POST":
		form = UpdatePersonnel(request.POST, request.FILES, instance=profile)

		if form.is_valid():
			instance = form.save()

			# For updating profile picture
			if 'raw_image' in request.FILES:
				save_profile_picture(instance)

			return redirect('profile-personnel', pk)
	else:
		form = UpdatePersonnel(instance=profile)

	context = {
		'prev'		: request.GET.get("prev", ""),
		'profile'	: profile,
		'form'		: form,
	}
	return render(request, "profiles/profile_update.html", context)


def profile_inmate_update(request, pk):
	profile = get_object_or_404(Inmate, pk=pk)

	if request.method == "POST":
		form = UpdateInmate(request.POST, request.FILES, instance=profile)

		if form.is_valid():
			instance = form.save()

			# For updating profile picture
			if 'raw_image' in request.FILES:
				save_profile_picture(instance)

			return redirect('profile-inmate', pk)
	else:
		form = UpdateInmate(instance=profile)

	context = {
		'prev'		: request.GET.get("prev", ""),
		'profile'	: profile,
		'form'		: form,
	}
	return render(request, "profiles/profile_update.html", context)


def profile_inmate_add(request):
	prev = request.GET.get("prev", "")

	if request.method == "POST":
		form = CreateInmate(request.POST, request.FILES)

		if form.is_valid():
			instance = form.save()
			if 'raw_image' in request.FILES:
				save_profile_picture(instance)
			return redirect(prev)
	else:
		form = CreateInmate()

	context = {
		'form' : form,
	}
	return render(request, "profiles/profile_add.html", context)


def profile_personnel_add(request):
	prev = request.GET.get("prev", "")

	if request.method == "POST":
		form = CreatePersonnel(request.POST, request.FILES)

		if form.is_valid():
			instance = form.save()
			if 'raw_image' in request.FILES:
				save_profile_picture(instance)
			return redirect(prev)
	else:
		form = CreatePersonnel()

	context = {
		'form' : form,
	}
	return render(request, "profiles/profile_add.html", context)


def profile_personnel_delete(request, pk):
	prev	= request.GET.get("prev", "")
	profile = get_object_or_404(Personnel, pk=pk)

	if request.method == "POST":
		profile.delete()

		if prev:
			return redirect(prev)
		else:
			return redirect('profiles-personnels')

	context = {
		'profile' : profile,
	}
	return render(request, "profiles/profile_delete_confirm.html", context)


def profile_inmate_delete(request, pk):
	prev	= request.GET.get("prev", "")
	profile = get_object_or_404(Inmate, pk=pk)

	if request.method == "POST":
		profile.delete()

		if prev:
			return redirect(prev)
		else:
			return redirect('profiles-inmates')

	context = {
		'profile' : profile,
	}
	return render(request, "profiles/profile_delete_confirm.html", context)


def profile_template_upload(request):
	if request.method == "POST":
		form = TemplateUploadForm(request.POST, request.FILES)

		if form.is_valid():
			form.save()
	else:
		form= TemplateUploadForm()

	context = {
		"form": form
	}
	return render(request, "profiles/profile_template_upload.html", context)

def profile_inmate_to_docx(request, pk):
	template	= Template.objects.get(template_name__icontains="inmate")
	profile		= Inmate.objects.get(pk=pk)

	fields = [
		'[[name]]',
		'[[age]]',
		'[[address]]',
		'[[civil_status]]',
		'[[date_arrested]]',
		'[[date_committed]]',
		'[[crime_violated]]'
	]

	if profile.date_arrested != None:
		date_arrested = datetime.fromisoformat(str(profile.date_arrested))
		date_arrested = date_arrested.strftime("%B %d, %Y, %I:%M %p")
	else:
		date_arrested = profile.date_arrested

	if profile.date_committed != None:
		date_committed = datetime.fromisoformat(str(profile.date_committed))
		date_committed = date_committed.strftime("%B %d, %Y, %I:%M %p")
	else:
		date_committed = profile.date_committed

	data = [
		get_full_name(profile),
		profile.age,
		profile.address,
		profile.civil_status.replace("_", " ").title(),
		date_arrested,
		date_committed,
		profile.crime_violated
	]

	doc = Document(template.template_file.path)

	for p in doc.paragraphs:
		if "raw_image" in p.text:
			p.clear()

			run = p.add_run()
			run.add_picture(
				profile.raw_image.path, 
				width=Inches(2), 
				height=Inches(2)
			)
			
			p.alignment = 1

	for table in doc.tables:
		for row in table.rows:
			cell = row.cells[1]
			# Check each paragraph in the second cell for placeholders
			for paragraph in cell.paragraphs:
				for field, value in zip(fields, data):
					if field in paragraph.text:
						paragraph.text = paragraph.text.replace(field, str(value))
	name		= str(get_full_name(profile, True))
	file_name	= f"{name}.docx"
	save_path	= outputs_path.joinpath(file_name)
	doc.save(save_path)

	with open(save_path, 'rb') as f:
		response = HttpResponse(
			f.read(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = f'attachment; filename="{file_name}.docx"'
	os.remove(save_path)
	return response


def profile_personnel_to_docx(request, pk):
	template	= Template.objects.get(template_name__icontains="personnel")
	profile		= Personnel.objects.get(pk=pk)

	fields = [
		'[[rank]]',
		'[[name]]',
		'[[age]]',
		'[[address]]',
		'[[civil_status]]',
		'[[date_assigned]]',
		'[[date_relieved]]',
		'[[designation]]'
	]

	if profile.date_assigned != None:
		date_assigned = datetime.fromisoformat(str(profile.date_assigned))
		date_assigned = date_assigned.strftime("%B %d, %Y, %I:%M %p")
	else:
		date_assigned = profile.date_assigned

	if profile.date_relieved != None:
		date_relieved = datetime.fromisoformat(str(profile.date_relieved))
		date_relieved = date_relieved.strftime("%B %d, %Y, %I:%M %p")
	else:
		date_relieved = profile.date_relieved

	data = [
		profile.rank.upper(),
		get_full_name(profile),
		profile.age,
		profile.address,
		profile.civil_status.replace("_", " ").title(),
		date_assigned,
		date_relieved,
		profile.designation
	]

	doc = Document(template.template_file.path)

	for p in doc.paragraphs:
		if "raw_image" in p.text:
			p.clear()

			run = p.add_run()
			run.add_picture(
				profile.raw_image.path, 
				width=Inches(2), 
				height=Inches(2)
			)
			
			p.alignment = 1

	for table in doc.tables:
		for row in table.rows:
			cell = row.cells[1]
			# Check each paragraph in the second cell for placeholders
			for paragraph in cell.paragraphs:
				for field, value in zip(fields, data):
					if field in paragraph.text:
						paragraph.text = paragraph.text.replace(field, str(value))

	name		= str(get_full_name(profile, True))
	file_name	= f"{name}.docx"
	save_path	= outputs_path.joinpath(file_name)
	doc.save(save_path)
	

	with open(save_path, 'rb') as f:
		response = HttpResponse(
			f.read(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = f'attachment; filename="{file_name}.docx"'
	os.remove(save_path)
	return response


def delete_all_inmate(request):
	prev 	= request.GET.get("prev", "")

	if request.method == "POST":
		Inmate.objects.all().delete()
		return redirect(prev)

	context = {
		'profile' : "all profiles",
	}
	return render(request, "profiles/profile_delete_confirm.html", context)


def delete_all_personnel(request):
	prev 	= request.GET.get("prev", "")

	if request.method == "POST":
		Personnel.objects.all().delete()
		return redirect(prev)

	context = {
		'profile' : "all profiles",
	}
	return render(request, "profiles/profile_delete_confirm.html", context)


 
def profile(request, p_type, pk):
	p_class =  Personnel if p_type == "personnel" else Inmate
	profile	=  get_object_or_404(p_class, pk=pk)

	context = {
		'profile': profile,
	}
	return render(request, "profiles/profile.html", context)
 
def profile_update(request, p_type, pk):
	p_class, update_form = [Personnel, UpdatePersonnel] if p_type == "personnel" else [Inmate, UpdateInmate]
	profile = get_object_or_404(p_class, pk=pk)

	if request.method == "POST":
		form = update_form(request.POST, request.FILES, instance=profile)

		if form.is_valid():
			instance = form.save()

			# For updating profile picture
			if 'raw_image' in request.FILES: save_profile_picture(instance)

			return redirect('profile', p_type, pk)
	else:
		form = update_form(instance=profile)

	context = {
		'prev'		: request.GET.get("prev", ""),
		'profile'	: profile,
		'form'		: form,
	}
	return render(request, "profiles/profile_update.html", context)