from django.http import HttpResponse
import string
from docx.shared import Inches
from docx import Document
from pathlib import Path
from django.conf import settings

def get_full_name(profile:object, m_initial:bool=True):
	full_name = profile.l_name

	if profile.f_name:
		full_name += f", {profile.f_name}"

	if profile.m_name and m_initial:
		full_name += f" {profile.m_name[0]}."
	elif profile.m_name:
		full_name += f" {profile.m_name}"

	if profile.suffix:
		full_name += f" {profile.suffix}"

	return full_name.title()

def format_text(text):
	symbols = string.punctuation.replace('_', ' ')
	translation_table = str.maketrans(symbols, '_' * len(symbols))

	text = text.replace(",", "")
	text = text.translate(translation_table)
	return text

def generate_docx(template_path, image_path, fields, data):
	doc = Document(template_path)

	# inserts the image
	for p in doc.paragraphs:
		if "raw_image" in p.text:
			p.clear()
			run = p.add_run()
			run.add_picture(
				image_path, 
				width=Inches(2), 
				height=Inches(2)
			)
			p.alignment = 1

	# fills in the table
	for table in doc.tables:
		for row in table.rows:
			cell = row.cells[1]
			for paragraph in cell.paragraphs:
				for field, value in zip(fields, data):
					if field in paragraph.text:
						for run in paragraph.runs:
							if field in run.text:
								run.text = run.text.replace(field, str(value))

							# if field in run.text:
							# 	run.text = f"{run.text.replace(field, '{{ {value} }}')}"
								run.font.name = 'Arial'

	file_name	= f"{format_text(data[-1])}.docx"
	if "[[rank]]" in fields:
		file_name	= f"{format_text(data[-5])}_{file_name}"

	save_path	= settings.MEDIA_ROOT.joinpath(file_name)
	doc.save(save_path)

	return [file_name, save_path]

def save_docx(file_name, save_path):
	with open(save_path, 'rb') as f:
		response = HttpResponse(
			f.read(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = f'attachment; filename="{file_name}"'
	
	save_path.unlink(save_path)

	return response